import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import font_manager as fm, rc
from scipy import stats
import os
import numpy as np
from matplotlib.ticker import FuncFormatter
from matplotlib.dates import DateFormatter
import seaborn as sns
import shutil
import matplotlib
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import win32com.client as win32
import pythoncom

# 폰트 크기 상수 정의
FONT_SIZE_LABEL = 14
FONT_SIZE_TICK = 12
FONT_SIZE_LEGEND = 12
FONT_SIZE_TITLE = 14

# 1. matplotlib 폰트 캐시 관련 코드 수정
try:
    cache_dir = matplotlib.get_cachedir()
    if os.path.exists(cache_dir):
        for file in os.listdir(cache_dir):
            file_path = os.path.join(cache_dir, file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(f'캐시 파일 삭제 중 오류 발생: {e}')
except Exception as e:
    print(f'캐시 디렉토리 처리 중 오류 발생: {e}')

# 2. 폰트 매니저 재생성
fm._load_fontmanager()

# 3. 폰트 설정
font_path = r"C:\Users\Y15599\Desktop\작업\폰트\윤고딕340.ttf"

try:
    font_prop = fm.FontProperties(fname=font_path)
    plt.rcParams['font.family'] = font_prop.get_name()
    plt.rcParams['axes.unicode_minus'] = False
    print(f"Font loaded successfully: {font_prop.get_name()}")
except Exception as e:
    print(f"Error loading font: {e}")
    
# 0 : 없음 / 

class BridgeAnalysis:
    def __init__(self, br_name):
        self.br_name = br_name
        # 인코딩 시도 함수
        def read_csv_with_encoding(file_path):
            try:
                return pd.read_csv(file_path, encoding='utf-8')
            except UnicodeDecodeError:
                try:
                    return pd.read_csv(file_path, encoding='cp949')
                except UnicodeDecodeError:
                    print(f"파일을 읽을 수 없습니다: {file_path}")
                    raise

        self.info = read_csv_with_encoding('01_channel_info/channel_info.csv')
        self.dynamic = read_csv_with_encoding(f'{self.br_name}/{self.br_name}_통계데이터.csv')
        try:
            self.static = read_csv_with_encoding(f'{self.br_name}/{self.br_name}_정적통계데이터.csv')
        except FileNotFoundError:
            print(f"{self.br_name}의 정적통계데이터가 없습니다.")
            self.static = pd.DataFrame(columns=['계측시간'])  # 빈 DataFrame 생성
        self.data_merge = pd.merge(self.dynamic, self.static, on='계측시간', how='outer')

        # date 변환을 초기화 단계에서 수행
        self.dynamic['계측시간'] = pd.to_datetime(self.dynamic['계측시간'])
        self.static['계측시간'] = pd.to_datetime(self.static['계측시간'])
        self.data_merge['계측시간'] = pd.to_datetime(self.data_merge['계측시간'])
        print(self.data_merge.info())

    def plot_time_history(self):
        """limit_type에 따라 그래프를 그리는 함수"""

        # info.csv에서 현재 교량(br_name)의 limit_type에 해당하는 channel_name 필터링
        filtered_info = self.info[self.info['br_name'] == self.br_name]
        channel_names = filtered_info['channel_name'].tolist()

        if not channel_names:
            print(f"br_name에 해당하는 채널이 없습니다: {self.br_name}")
            return
        print(channel_names)

        data = self.data_merge

        for idx, col in enumerate(channel_names):
            avg_col = f"{col}_AVG"
            value_col = f"{col}_VALUE"
            
            # AVG 컬럼이 없으면 VALUE 컬럼 사용
            if avg_col not in data.columns:
                if value_col in data.columns:
                    avg_col = value_col
                else:
                    print(f"{avg_col} 또는 {value_col} 컬럼이 올해 데이터에 없습니다. 스킵합니다.")
                    continue

            # limit_type에 따른 상한선 또는 하한선 가져오기
            channel_info = filtered_info[filtered_info['channel_name'] == col].iloc[0]
            limit_type = channel_info['limit_type']
            b = channel_info.get('b')
            h = channel_info.get('h')
            up_limit1 = channel_info.get('up_limit1', None)
            up_limit2 = channel_info.get('up_limit2', None)
            low_limit2 = channel_info.get('low_limit2', None)
            low_limit1 = channel_info.get('low_limit1', None)
            limit1 = channel_info.get('limit1', None)
            limit2 = channel_info.get('limit2', None)
            limit3 = channel_info.get('limit3', None)
            limit4 = channel_info.get('limit4', None)
            significant_figure = channel_info.get('significant_figure', None)
            correl = channel_info.get('correl',None)
            legendloc = channel_info.get('legendloc', None)
            bbox_x = channel_info.get('bbox_x', None)
            bbox_y = channel_info.get('bbox_y', None)
            eqk1 = channel_info.get('eqk1', None)
            eqk2 = channel_info.get('eqk2', None)

            # ylabel 구성: 채널종류 + (단위)
            ylabel = f"{channel_info['채널종류']} ({channel_info['unit']})"
            
            def format_y_ticks(value, _):
                if value < 0: 
                    return f'-{abs(value):,.{significant_figure}f}'
                else:
                    return f'{value:,.{significant_figure}f}'

            # 그래프 생성
            fig, ax = plt.subplots(figsize=(b, h))

            if limit_type == 0 or limit_type == 6:  # limit_type 0 시간이력그래프
                plot_data = data[['계측시간', avg_col]].dropna()
                if plot_data.empty:
                    print(f"{avg_col} 컬럼에 유효한 데이터가 없습니다. 스킵합니다.")
                    continue
                ax.plot(plot_data['계측시간'], plot_data[avg_col], label=col, alpha=0.8, color='darkorange')

            elif limit_type == 1:  # limit_type 1: 상한선이 있는 시간이력그래프
                plot_data = data[['계측시간', avg_col]].dropna()
                if plot_data.empty:
                    print(f"{avg_col} 컬럼에 유효한 데이터가 없습니다. 스킵합니다.")
                    continue
                ax.axhline(y=up_limit1, color='red', linestyle='-', linewidth=2, label=f'알람상한: {round(up_limit1,significant_figure)}')
                ax.plot(plot_data['계측시간'], plot_data[avg_col], label=col, alpha=0.8, color='darkorange')

            elif limit_type == 2:  # limit_type 2: 상한선이 2개 있는 시간이력그래프
                plot_data = data[['계측시간', avg_col]].dropna()
                if plot_data.empty:
                    print(f"{avg_col} 컬럼에 유효한 데이터가 없습니다. 스킵합니다.")
                    continue
                ax.axhline(y=up_limit2, color='red', linestyle='-', linewidth=2, label=f'관리상한: {round(up_limit2,significant_figure)}')
                ax.axhline(y=up_limit1, color='red', linestyle='--', linewidth=2, label=f'알람상한: {round(up_limit1,significant_figure)}')
                ax.plot(plot_data['계측시간'], plot_data[avg_col], label=col, alpha=0.8, color='darkorange')

            elif limit_type == 4:  # limit_type 4: 상한선, 하한선이 각 2개 있는 시간이력그래프
                plot_data = data[['계측시간', avg_col]].dropna()
                if plot_data.empty:
                    print(f"{avg_col} 컬럼에 유효한 데이터가 없습니다. 스킵합니다.")
                    continue
                ax.axhline(y=up_limit2, color='red', linestyle='-', linewidth=2, label=f'관리상한: {round(up_limit2,significant_figure)}')
                ax.axhline(y=up_limit1, color='red', linestyle='--', linewidth=2, label=f'알림상한: {round(up_limit1,significant_figure)}')
                ax.plot(plot_data['계측시간'], plot_data[avg_col], label=col, alpha=0.8, color='darkorange')
                ax.axhline(y=low_limit1, color='blue', linestyle='--', linewidth=2, label=f'알림하한: {round(low_limit1,significant_figure)}')
                ax.axhline(y=low_limit2, color='blue', linestyle='-', linewidth=2, label=f'관리하한: {round(low_limit2,significant_figure)}')

            elif limit_type == 5:  # limit_type 5: 상한선 1개 있는 시간이력 표준편차그래프
                sdt_col = f"{col}_STD"
                if sdt_col not in data.columns:
                    print(f"{sdt_col} 컬럼이 올해 데이터에 없습니다. 스킵합니다.")
                    continue
                
                plot_data = data[['계측시간', sdt_col]].dropna()
                if plot_data.empty:
                    print(f"{sdt_col} 컬럼에 유효한 데이터가 없습니다. 스킵합니다.")
                    continue
                    
                '''ax.axhline(y=50, color='red', linestyle='--', linewidth=2, label=f'알림상한: 50 gal')'''
                ax.plot(plot_data['계측시간'], plot_data[sdt_col], label=col, alpha=0.8, color='darkorange')

            else:
                pass

            # 그래프 설정
            ax.set_ylabel(ylabel, fontsize=FONT_SIZE_LABEL)
            ax.set_title(f'{self.br_name} {col}', fontsize=FONT_SIZE_TITLE)
            ax.legend(loc=legendloc, bbox_to_anchor=(bbox_x, bbox_y), prop=font_prop, fontsize=FONT_SIZE_LEGEND, ncol=5)
            ax.yaxis.set_major_formatter(FuncFormatter(format_y_ticks))
            
            # x축 날짜 포맷 설정
            ax.xaxis.set_major_formatter(DateFormatter('%m-%d'))
            #plt.xticks(rotation=45, ha='right')
            ax.set_xlim(plot_data['계측시간'].min(), plot_data['계측시간'].max())
            ax.tick_params(axis='both', labelsize=FONT_SIZE_TICK)
            ax.grid(True)

            # 그래프 저장
            plt.savefig(f'{self.br_name}/{col}.png', dpi=300, bbox_inches='tight') 
            plt.close()
            print(f'{idx + 1}/{len(channel_names)}: {col} 그래프 생성 완료')

    def plot_scatter(self):
        # info.csv에서 현재 교량(br_name)의 limit_type에 해당하는 channel_name 필터링
        filtered_info1 = self.info[(self.info['br_name'] == self.br_name)&(self.info['limit_type'].isin([6, 12, 13, 14]))]
        channel_names = filtered_info1['channel_name'].tolist()
        print(channel_names)

        if not channel_names:
            print(f"br_name에 해당하는 채널이 없습니다: {self.br_name}")
            return

        data = self.data_merge

        for idx, col in enumerate(channel_names):
            avg_col = f"{col}_AVG"
            value_col = f"{col}_VALUE"
            
            # AVG 컬럼이 없으면 VALUE 컬럼 사용
            if avg_col not in data.columns:
                if value_col in data.columns:
                    avg_col = value_col
                else:
                    print(f"{avg_col} 또는 {value_col} 컬럼이 올해 데이터에 없습니다. 스킵합니다.")
                    continue

            # limit_type에 따른 상관관계정보 가져오기
            channel_info1 = filtered_info1[filtered_info1['channel_name'] == col].iloc[0]
            limit_type = channel_info1['limit_type']
            b = channel_info1.get('b')
            h = channel_info1.get('h')
            correl = channel_info1['correl']
            avg_correl = f"{correl}_AVG"
            value_correl = f"{correl}_VALUE"

            # 상관관계 컬럼 확인
            if avg_correl not in data.columns:
                if value_correl in data.columns:
                    avg_correl = value_correl
                else:
                    print(f"{avg_correl} 또는 {value_correl} 컬럼이 올해 데이터에 없습니다. 스킵합니다.")
                    continue

            # ylabel 구성: 채널종류 + (단위)
            ylabel = f"{channel_info1['채널종류']} ({channel_info1['unit']})"

            def format_y_ticks(value, _):
                if value < 0:
                    return f'-{abs(value):,.0f}'
                return f'{value:,.0f}'
            
            print(idx, col, correl)
            
            if limit_type == 6:  # limit_type 6: 산점도
                # matplotlib 스타일 설정
                plt.style.use('default')
                matplotlib.rcParams['axes.unicode_minus'] = False  # 음수 기호 표시 설정
                
                fig, ax = plt.subplots(figsize=(b, h))

                # NaN 제거
                valid_data = self.data_merge[[avg_col, avg_correl]].dropna()

                if valid_data.empty:
                    print(f"데이터가 부족하여 {avg_col} 산점도를 생성할 수 없습니다.")
                    continue

                x_data = valid_data[avg_correl]
                y_data = valid_data[avg_col]

                ax.scatter(x_data, y_data)

                # 회귀 분석 수행
                slope, intercept, r_value, p_value, std_err = stats.linregress(x_data, y_data)

                # 회귀선 추가
                ax.plot(x_data, intercept + slope * x_data, color='red', 
                        label=f'y = {slope:.2f}x + {intercept:.2f}\n$R^2$ = {r_value**2:.2f}')

                # y축 포맷터 설정
                def format_y_ticks(value, _):
                    if value < 0:
                        return f'-{abs(value):,.0f}'
                    return f'{value:,.0f}'
                
                # y축, x축 포맷터 적용
                formatter = FuncFormatter(format_y_ticks)
                ax.yaxis.set_major_formatter(formatter)
                ax.xaxis.set_major_formatter(formatter)
                
                # y축 범위 설정
                y_min = min(y_data.min(), 0)  # 음수 값이 있다면 포함
                y_max = max(y_data.max(), 0)  # 양수 값이 있다면 포함
                ax.set_ylim(y_min - 0.1 * abs(y_min), y_max + 0.1 * abs(y_max))

                # y축, x축 눈금 설정
                ax.tick_params(axis='y', which='major', labelsize=12)
                ax.tick_params(axis='x', which='major', labelsize=12)
                
                ax.set_xlabel(f'{channel_info1["correl"]} (ºC)', fontsize=14, fontweight='bold', font=font_prop)
                ax.set_ylabel(ylabel, fontsize=14, fontweight='bold', font=font_prop)
                ax.legend(loc='best', prop=font_prop)
                ax.grid(True)
                fig.tight_layout()

                # 결과 저장
                output_dir = f'{self.br_name}'
                os.makedirs(output_dir, exist_ok=True)
                plt.savefig(f'{output_dir}/{col}_.jpg')
                plt.close(fig)
                print(f'{idx + 1}/{len(channel_names)}: {col} 그래프 생성 완료')
            else:
                pass

    def generate_summary_report(self):
        """월간 요약 보고서를 생성하는 함수"""
        # info.csv에서 현재 교량(br_name)의 limit_type에 해당하는 channel_name 필터링
        filtered_info = self.info[self.info['br_name'] == self.br_name]
        channel_names = filtered_info['channel_name'].tolist()

        if not channel_names:
            print(f"br_name에 해당하는 채널이 없습니다: {self.br_name}")
            return

        data = self.data_merge

        # HTML 보고서 시작
        html_content = f"""
        <!DOCTYPE html>
        <html lang="ko">
        <head>
            <meta charset="UTF-8">
            <title>{self.br_name} 계측 데이터 요약 보고서</title>
            <style>
                body {{
                    font-family: 'Malgun Gothic', sans-serif;
                    margin: 40px;
                    color: #333;
                }}
                .container {{
                    max-width: 800px;
                    margin: 0 auto;
                }}
                h1 {{
                    text-align: center;
                    color: #2c3e50;
                    margin-bottom: 30px;
                }}
                .section {{
                    margin-bottom: 30px;
                    padding: 20px;
                    background-color: #fff;
                    border-radius: 5px;
                    box-shadow: 0 2px 5px rgba(0,0,0,0.1);
                }}
                .section-title {{
                    font-size: 1.2em;
                    color: #2c3e50;
                    margin-bottom: 15px;
                    padding-bottom: 10px;
                    border-bottom: 2px solid #eee;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                }}
                th, td {{
                    padding: 12px;
                    text-align: left;
                    border-bottom: 1px solid #ddd;
                }}
                th {{
                    background-color: #f5f6fa;
                }}
                .progress-bar {{
                    width: 100%;
                    background-color: #f0f0f0;
                    padding: 3px;
                    border-radius: 3px;
                    box-shadow: inset 0 1px 3px rgba(0, 0, 0, .2);
                }}
                .progress-bar-fill {{
                    display: block;
                    height: 22px;
                    background-color: #659DBD;
                    border-radius: 3px;
                    transition: width 500ms ease-in-out;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>{self.br_name} 계측 데이터 요약 보고서</h1>
                
                <!-- 1. 데이터 수집 정보 -->
                <div class="section">
                    <div class="section-title">1. 데이터 수집 정보</div>
                    <table>
                        <tr>
                            <th>항목</th>
                            <th>내용</th>
                        </tr>
                        <tr>
                            <td>데이터 수집 기간</td>
                            <td>{data['계측시간'].min().strftime('%Y-%m-%d')} ~ {data['계측시간'].max().strftime('%Y-%m-%d')}</td>
                        </tr>
                        <tr>
                            <td>수집 주기</td>
                            <td>{pd.Timedelta(data['계측시간'].diff().mode().iloc[0]).total_seconds() / 60:.0f}분</td>
                        </tr>
                        <tr>
                            <td>센서 종류</td>
                            <td>{", ".join(filtered_info['센서종류'].unique())}</td>
                        </tr>
                    </table>
                </div>

                <!-- 2. 센서 종류별 수량 -->
                <div class="section">
                    <div class="section-title">2. 센서 종류별 수량</div>
                    <table>
                        <tr>
                            <th>센서 종류</th>
                            <th>수량</th>
                        </tr>
        """

        # 센서 종류별 수량 계산
        sensor_counts = filtered_info['센서종류'].value_counts()
        for sensor_type, count in sensor_counts.items():
            html_content += f"<tr><td>{sensor_type}</td><td>{count}</td></tr>"
        
        # 합계 행 추가
        html_content += f"""
                        <tr style="font-weight: bold;">
                            <td>합계</td>
                            <td>{sensor_counts.sum()}</td>
                        </tr>
                    </table>
                </div>

                <!-- 3. 센서별 데이터 수신율 -->
                <div class="section">
                    <div class="section-title">3. 센서별 데이터 수신율</div>
                    <table>
                        <tr>
                            <th>센서명</th>
                            <th>수신율</th>
                        </tr>
        """

        # 데이터 수신율 계산
        total_expected_records = len(pd.date_range(start=data['계측시간'].min(), 
                                                 end=data['계측시간'].max(), 
                                                 freq=f"{pd.Timedelta(data['계측시간'].diff().mode().iloc[0]).total_seconds() / 60:.0f}min"))
        
        for col in channel_names:
            avg_col = f"{col}_AVG"
            value_col = f"{col}_VALUE"
            
            # AVG 컬럼이 없으면 VALUE 컬럼 사용
            if avg_col not in data.columns:
                if value_col in data.columns:
                    avg_col = value_col
                else:
                    continue

            actual_records = data[avg_col].count()
            reception_rate = (actual_records / total_expected_records) * 100
            html_content += f"""
                        <tr>
                            <td>{col}</td>
                            <td>
                                <div class="progress-bar">
                                    <span class="progress-bar-fill" style="width: {reception_rate}%"></span>
                                </div>
                                {actual_records}/{total_expected_records} ({reception_rate:.1f}%)
                            </td>
                        </tr>
            """

        html_content += """
                    </table>
                </div>

                <!-- 4. 이상치 분석 -->
                <div class="section">
                    <div class="section-title">4. 이상치 분석</div>
                    <table>
                        <tr>
                            <th>센서명</th>
                            <th>이상치 수량</th>
                            <th>최솟값</th>
                            <th>평균값</th>
                            <th>최댓값</th>
                            <th>범위(최대-최소)</th>
                        </tr>
        """

        # 이상치 분석
        for col in channel_names:
            avg_col = f"{col}_AVG"
            if avg_col in data.columns:
                plot_data = data[['계측시간', avg_col]].dropna()
                if not plot_data.empty:
                    Q1 = plot_data[avg_col].quantile(0.25)
                    Q3 = plot_data[avg_col].quantile(0.75)
                    IQR = Q3 - Q1
                    outliers = plot_data[avg_col][(plot_data[avg_col] < (Q1 - 1.5 * IQR)) | (plot_data[avg_col] > (Q3 + 1.5 * IQR))]
                    
                    # significant_figure 가져오기
                    significant_figure = filtered_info[filtered_info['channel_name'] == col]['significant_figure'].iloc[0]
                    
                    html_content += f"""
                        <tr>
                            <td>{col}</td>
                            <td>{len(outliers)}</td>
                            <td>{plot_data[avg_col].min():.{significant_figure}f}</td>
                            <td>{plot_data[avg_col].mean():.{significant_figure}f}</td>
                            <td>{plot_data[avg_col].max():.{significant_figure}f}</td>
                            <td>{plot_data[avg_col].max() - plot_data[avg_col].min():.{significant_figure}f}</td>
                        </tr>
                    """

        html_content += """
                    </table>
                </div>

                <!-- 5. 관리기준 초과 여부 -->
                <div class="section">
                    <div class="section-title">5. 관리기준 초과 여부</div>
                    <table>
                        <tr>
                            <th>센서명</th>
                            <th>최댓값</th>
                            <th>최솟값</th>
                            <th>관리상한</th>
                            <th>관리하한</th>
                            <th>초과여부</th>
                            <th>비고</th>
                        </tr>
        """

        # 관리기준 초과 여부 분석
        for col in channel_names:
            avg_col = f"{col}_AVG"
            if avg_col in data.columns:
                channel_info = filtered_info[filtered_info['channel_name'] == col].iloc[0]
                limit_type = channel_info['limit_type']
                plot_data = data[['계측시간', avg_col]].dropna()
                
                if not plot_data.empty:
                    max_value = plot_data[avg_col].max()
                    min_value = plot_data[avg_col].min()
                    significant_figure = channel_info['significant_figure']
                    
                    # limit_type에 따른 관리기준 설정
                    if limit_type == 2:
                        up_limit = channel_info.get('up_limit2')  # 관리상한
                        low_limit = None  # 하한 없음
                    elif limit_type == 4:
                        up_limit = channel_info.get('up_limit2')  # 관리상한
                        low_limit = channel_info.get('low_limit2')  # 관리하한
                    else:  # limit_type이 0 또는 6인 경우
                        up_limit = None
                        low_limit = None
                    
                    # 초과 여부 확인
                    is_exceeded = False
                    note = ""
                    
                    if limit_type in [2, 4]:
                        if up_limit is not None and max_value > up_limit:
                            is_exceeded = True
                            note += f"최댓값이 관리상한({up_limit}) 초과\n"
                        
                        if low_limit is not None and min_value < low_limit:
                            is_exceeded = True
                            note += f"최솟값이 관리하한({low_limit}) 미만\n"
                    else:  # limit_type이 0 또는 6인 경우
                        note = "관리기준 없음"
                    
                    # HTML 스타일 설정
                    style = 'color: red;' if is_exceeded else ''
                    
                    html_content += f"""
                        <tr style="{style}">
                            <td>{col}</td>
                            <td>{max_value:.{significant_figure}f}</td>
                            <td>{min_value:.{significant_figure}f}</td>
                            <td>{up_limit if up_limit is not None else "-"}</td>
                            <td>{low_limit if low_limit is not None else "-"}</td>
                            <td>{"초과" if is_exceeded else "정상"}</td>
                            <td>{note}</td>
                        </tr>
                    """

        html_content += """
                    </table>
                </div>
            </div>
        </body>
        </html>
        """

        # HTML 파일 저장
        output_dir = f'{self.br_name}'
        os.makedirs(output_dir, exist_ok=True)
        with open(f'{output_dir}/{self.br_name}_요약보고서.html', 'w', encoding='utf-8') as f:
            f.write(html_content)

        # 통계 출력
        for idx, col in enumerate(channel_names):
            avg_col = f"{col}_AVG"
            if avg_col not in data.columns:
                print(f"{avg_col} 컬럼이 올해 데이터에 없습니다. 스킵합니다.")
                continue

            # limit_type에 따른 상한선 또는 하한선 가져오기
            channel_info = filtered_info[filtered_info['channel_name'] == col].iloc[0]
            limit_type = channel_info['limit_type']
            up_limit1 = channel_info.get('up_limit1', None)
            up_limit2 = channel_info.get('up_limit2', None)
            low_limit2 = channel_info.get('low_limit2', None)
            low_limit1 = channel_info.get('low_limit1', None)
            limit1 = channel_info.get('limit1', None)
            limit2 = channel_info.get('limit2', None)
            limit3 = channel_info.get('limit3', None)
            limit4 = channel_info.get('limit4', None)
            significant_figure = channel_info.get('significant_figure', None)

            # 날짜와 값이 모두 유효한 데이터만 추출
            plot_data = data[['계측시간', avg_col]].dropna()
            if plot_data.empty:
                print(f"{avg_col} 컬럼에 유효한 데이터가 없습니다. 스킵합니다.")
                continue

            # 통계 계산
            mean_value = plot_data[avg_col].mean()
            std_value = plot_data[avg_col].std()
            max_value = plot_data[avg_col].max()
            min_value = plot_data[avg_col].min()

            # limit_type에 따른 알람 체크
            if limit_type == 1:  # limit_type 1: 상한선이 있는 시간이력그래프
                if max_value > up_limit1:
                    print(f"{col} 알람상한 초과: {max_value:.{significant_figure}f} > {up_limit1:.{significant_figure}f}")

            elif limit_type == 2:  # limit_type 2: 상한선이 2개 있는 시간이력그래프
                if max_value > up_limit2:
                    print(f"{col} 관리상한 초과: {max_value:.{significant_figure}f} > {up_limit2:.{significant_figure}f}")
                elif max_value > up_limit1:
                    print(f"{col} 알람상한 초과: {max_value:.{significant_figure}f} > {up_limit1:.{significant_figure}f}")

            elif limit_type == 4:  # limit_type 4: 상한선, 하한선이 각 2개 있는 시간이력그래프
                if max_value > up_limit2:
                    print(f"{col} 관리상한 초과: {max_value:.{significant_figure}f} > {up_limit2:.{significant_figure}f}")
                elif max_value > up_limit1:
                    print(f"{col} 알림상한 초과: {max_value:.{significant_figure}f} > {up_limit1:.{significant_figure}f}")
                if min_value < low_limit2:
                    print(f"{col} 관리하한 초과: {min_value:.{significant_figure}f} < {low_limit2:.{significant_figure}f}")
                elif min_value < low_limit1:
                    print(f"{col} 알림하한 초과: {min_value:.{significant_figure}f} < {low_limit1:.{significant_figure}f}")

            # 통계 출력
            print(f"\n{col} 통계:")
            print(f"평균: {mean_value:.{significant_figure}f}")
            print(f"표준편차: {std_value:.{significant_figure}f}")
            print(f"최대값: {max_value:.{significant_figure}f}")
            print(f"최소값: {min_value:.{significant_figure}f}")

    def generate_summary_report_excel(self):
        """교량 계측 데이터 요약 보고서를 Excel 형식으로 생성하는 함수"""
        
        # 1. 데이터 수집 정보
        start_date = self.data_merge['계측시간'].min()
        end_date = self.data_merge['계측시간'].max()
        sampling_period = pd.Timedelta(self.data_merge['계측시간'].diff().mode().iloc[0]).total_seconds() / 60  # 분 단위로 변환
        
        # 2. 계측기 종류 및 수량
        sensor_info = self.info[self.info['br_name'] == self.br_name]
        sensor_counts = sensor_info['센서종류'].value_counts()
        
        # 3. 센서별 데이터 수신율 계산 및 연속 결측 기간 확인
        total_expected_records = len(pd.date_range(start=start_date, end=end_date, freq=f'{sampling_period}min'))
        reception_rates = {}
        missing_periods = {}  # 연속 결측 기간을 저장할 딕셔너리
        
        for col in sensor_info['channel_name']:
            avg_col = f"{col}_AVG"
            value_col = f"{col}_VALUE"
            
            # AVG 컬럼이 없으면 VALUE 컬럼 사용
            if avg_col not in self.data_merge.columns:
                if value_col in self.data_merge.columns:
                    avg_col = value_col
                else:
                    continue

            # 데이터 수신율 계산
            actual_records = self.data_merge[avg_col].count()
            total_missing = total_expected_records - actual_records
            reception_rates[col] = {
                'actual': actual_records,
                'total': total_expected_records,
                'rate': (actual_records / total_expected_records) * 100,
                'missing': total_missing
            }
            
            # 연속 결측 기간 확인
            missing_data = self.data_merge[avg_col].isna()
            missing_periods[col] = []
            
            if missing_data.any():
                # 결측(True)인 구간만 연속 그룹핑
                group_id = (missing_data != missing_data.shift()).cumsum()
                true_groups = group_id[missing_data]
                for gid in true_groups.unique():
                    group_idx = (group_id == gid) & missing_data
                    group_data = self.data_merge[group_idx]
                    if len(group_data) >= 60:  # 60개 이상 연속된 결측
                        start_time = group_data['계측시간'].min()
                        end_time = group_data['계측시간'].max()
                        missing_periods[col].append({
                            'start': start_time,
                            'end': end_time,
                            'duration': len(group_data)
                        })
        
        # 4. 이상치 확인 및 통계 정보
        outliers_summary = {}
        stats_summary = {}
        for col in sensor_info['channel_name']:
            avg_col = f"{col}_AVG"
            value_col = f"{col}_VALUE"
            
            # AVG 컬럼이 없으면 VALUE 컬럼 사용
            if avg_col not in self.data_merge.columns:
                if value_col in self.data_merge.columns:
                    avg_col = value_col
                else:
                    continue

            data = self.data_merge[avg_col].dropna()
            if len(data) > 0:
                Q1 = data.quantile(0.25)
                Q3 = data.quantile(0.75)
                IQR = Q3 - Q1
                outliers = data[(data < (Q1 - 1.5 * IQR)) | (data > (Q3 + 1.5 * IQR))]
                outliers_summary[col] = len(outliers)
                
                # significant_figure 가져오기
                significant_figure = sensor_info[sensor_info['channel_name'] == col]['significant_figure'].iloc[0]
                
                # 통계 정보 계산
                stats_summary[col] = {
                    'min': data.min(),
                    'mean': data.mean(),
                    'max': data.max(),
                    'range': data.max() - data.min(),
                    'significant_figure': significant_figure
                }

        # Excel 워크북 생성
        wb = Workbook()
        
        # 스타일 설정
        header_fill = PatternFill(start_color='E6E6E6', end_color='E6E6E6', fill_type='solid')
        header_font = Font(bold=True)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        center_aligned = Alignment(horizontal='center', vertical='center')
        
        # 1. 데이터 수집 정보 시트
        ws1 = wb.active
        ws1.title = "데이터 수집 정보"
        
        ws1['A1'] = f"{self.br_name} 계측 데이터 요약 보고서"
        ws1['A1'].font = Font(size=14, bold=True)
        ws1.merge_cells('A1:B1')
        
        headers = ['항목', '내용']
        for col, header in enumerate(headers, 1):
            cell = ws1.cell(row=2, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = center_aligned
        
        data = [
            ['데이터 수집 기간', f"{start_date.strftime('%Y-%m-%d')} ~ {end_date.strftime('%Y-%m-%d')}"],
            ['수집 주기', f"{sampling_period:.0f}분"],
            ['센서 종류', ", ".join(sensor_info["센서종류"].unique())]
        ]
        
        for row_idx, row_data in enumerate(data, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws1.cell(row=row_idx, column=col_idx)
                cell.value = value
                cell.border = border
                cell.alignment = center_aligned
        
        # 2. 센서 종류별 수량 시트
        ws2 = wb.create_sheet("센서 종류별 수량")
        
        ws2['A1'] = "센서 종류별 수량"
        ws2['A1'].font = Font(size=14, bold=True)
        ws2.merge_cells('A1:B1')
        
        headers = ['센서 종류', '수량']
        for col, header in enumerate(headers, 1):
            cell = ws2.cell(row=2, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = center_aligned
        
        for row_idx, (sensor_type, count) in enumerate(sensor_counts.items(), 3):
            ws2.cell(row=row_idx, column=1, value=sensor_type).border = border
            ws2.cell(row=row_idx, column=2, value=count).border = border
        
        # 합계 행 추가
        total_row = len(sensor_counts) + 3
        ws2.cell(row=total_row, column=1, value="합계").font = Font(bold=True)
        ws2.cell(row=total_row, column=2, value=sensor_counts.sum()).font = Font(bold=True)
        
        # 3. 센서별 데이터 수신율 시트
        ws3 = wb.create_sheet("데이터 수신율")
        
        ws3['A1'] = "센서별 데이터 수신율"
        ws3['A1'].font = Font(size=14, bold=True)
        ws3.merge_cells('A1:C1')
        
        headers = ['센서명', '수신데이터수(개)', '전체데이터수(개)', '수신율(%)', '전체결측수']
        for col, header in enumerate(headers, 1):
            cell = ws3.cell(row=2, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = center_aligned
            
        # 각 센서별 데이터 입력
        for row_idx, (channel, rate) in enumerate(reception_rates.items(), 3):
            ws3.cell(row=row_idx, column=1, value=channel).border = border  # 센서명
            ws3.cell(row=row_idx, column=2, value=rate['actual']).border = border  # 수신데이터수
            ws3.cell(row=row_idx, column=3, value=rate['total']).border = border  # 전체데이터수 
            ws3.cell(row=row_idx, column=4, value=f"{rate['rate']:.1f}").border = border  # 수신율
            ws3.cell(row=row_idx, column=5, value=rate['missing']).border = border  # 전체결측수
        
        # 4. 이상치 분석 시트
        ws4 = wb.create_sheet("이상치 분석")
        
        ws4['A1'] = "이상치 분석"
        ws4['A1'].font = Font(size=14, bold=True)
        ws4.merge_cells('A1:F1')
        
        headers = ['센서명', '이상치 수량', '최솟값', '평균값', '최댓값', '범위(최대-최소)']
        for col, header in enumerate(headers, 1):
            cell = ws4.cell(row=2, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = center_aligned
        
        for row_idx, (channel, count) in enumerate(outliers_summary.items(), 3):
            ws4.cell(row=row_idx, column=1, value=channel).border = border
            ws4.cell(row=row_idx, column=2, value=count).border = border
            ws4.cell(row=row_idx, column=3, value=f"{stats_summary[channel]['min']:.{stats_summary[channel]['significant_figure']}f}").border = border
            ws4.cell(row=row_idx, column=4, value=f"{stats_summary[channel]['mean']:.{stats_summary[channel]['significant_figure']}f}").border = border
            ws4.cell(row=row_idx, column=5, value=f"{stats_summary[channel]['max']:.{stats_summary[channel]['significant_figure']}f}").border = border
            ws4.cell(row=row_idx, column=6, value=f"{stats_summary[channel]['range']:.{stats_summary[channel]['significant_figure']}f}").border = border
        
        # 5. 관리기준 초과 여부 시트
        ws5 = wb.create_sheet("관리기준 초과 여부")
        
        ws5['A1'] = "관리기준 초과 여부"
        ws5['A1'].font = Font(size=14, bold=True)
        ws5.merge_cells('A1:G1')
        
        headers = ['센서명', '최댓값', '최솟값', '관리상한', '관리하한', '초과여부', '비고']
        for col, header in enumerate(headers, 1):
            cell = ws5.cell(row=2, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = center_aligned

        row_idx = 3
        for channel in sensor_info['channel_name']:
            avg_col = f"{channel}_AVG"
            value_col = f"{channel}_VALUE"
            
            # AVG 컬럼이 없으면 VALUE 컬럼 사용
            if avg_col not in self.data_merge.columns:
                if value_col in self.data_merge.columns:
                    avg_col = value_col
                else:
                    continue

            channel_info = sensor_info[sensor_info['channel_name'] == channel].iloc[0]
            limit_type = channel_info['limit_type']
            data = self.data_merge[avg_col].dropna()
            
            if len(data) > 0:
                max_value = data.max()
                min_value = data.min()
                
                # limit_type에 따른 관리기준 설정
                if limit_type == 2:
                    up_limit = channel_info.get('up_limit2')  # 관리상한
                    low_limit = None  # 하한 없음
                elif limit_type == 4:
                    up_limit = channel_info.get('up_limit2')  # 관리상한
                    low_limit = channel_info.get('low_limit2')  # 관리하한
                else:  # limit_type이 0 또는 6인 경우
                    up_limit = None
                    low_limit = None
                
                # 초과 여부 확인 (limit_type이 2 또는 4인 경우에만)
                is_exceeded = False
                note = ""
                
                if limit_type in [2, 4]:
                    if up_limit is not None and max_value > up_limit:
                        is_exceeded = True
                        note += f"최댓값이 관리상한({up_limit}) 초과\n"
                    
                    if low_limit is not None and min_value < low_limit:
                        is_exceeded = True
                        note += f"최솟값이 관리하한({low_limit}) 미만\n"
                else:  # limit_type이 0 또는 6인 경우
                    note = "관리기준 없음"
                
                # 데이터 입력
                ws5.cell(row=row_idx, column=1, value=channel).border = border
                ws5.cell(row=row_idx, column=2, value=f"{max_value:.{channel_info['significant_figure']}f}").border = border
                ws5.cell(row=row_idx, column=3, value=f"{min_value:.{channel_info['significant_figure']}f}").border = border
                ws5.cell(row=row_idx, column=4, value=str(up_limit) if up_limit is not None else "-").border = border
                ws5.cell(row=row_idx, column=5, value=str(low_limit) if low_limit is not None else "-").border = border
                ws5.cell(row=row_idx, column=6, value="초과" if is_exceeded else "정상").border = border
                ws5.cell(row=row_idx, column=7, value=note).border = border
                
                # 초과된 경우 빨간색으로 표시
                if is_exceeded:
                    for col in range(1, 8):
                        ws5.cell(row=row_idx, column=col).font = Font(color="FF0000")
                
                row_idx += 1
        
        # 열 너비 자동 조정
        for ws in [ws1, ws2, ws3, ws4, ws5]:
            for column in ws.columns:
                max_length = 0
                column = [cell for cell in column]
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width
        
        # Excel 파일 저장
        os.makedirs(self.br_name, exist_ok=True)
        wb.save(f'{self.br_name}/{self.br_name}_요약보고서.xlsx')
        
        print(f'{self.br_name} 요약 보고서가 Excel 형식으로 생성되었습니다.')

    def generate_summary_report_word(self):
        """엑셀 요약 보고서를 워드 문서로 변환하는 함수"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            import pandas as pd
            
            # 엑셀 파일 읽기
            excel_file = f'{self.br_name}/{self.br_name}_요약보고서.xlsx'
            if not os.path.exists(excel_file):
                print("엑셀 요약 보고서가 존재하지 않습니다. 먼저 엑셀 보고서를 생성해주세요.")
                return
            
            # 워드 문서 생성
            doc = Document()
            
            # 기본 폰트 설정
            doc.styles['Normal'].font.name = '맑은 고딕'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 제목 추가
            title = doc.add_heading(f'{self.br_name} 계측 데이터 요약 보고서', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 엑셀 파일의 각 시트 읽기
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            # 각 시트의 데이터를 워드 문서에 추가
            for sheet_name, df in excel_data.items():
                # 시트 제목 추가
                doc.add_heading(sheet_name, level=1)
                
                # 테이블 생성 (헤더 행 제외)
                table = doc.add_table(rows=len(df), cols=len(df.columns))
                table.style = 'Table Grid'
                
                # 데이터 추가
                for i, row in df.iterrows():
                    for j, value in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = str(value)
                        # 데이터 셀 스타일 설정
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                
                # 시트 간 간격 추가
                doc.add_paragraph()
            
            # 그래프 섹션 추가
            doc.add_heading("시간이력 그래프", level=1)
            
            # 그래프 파일 목록 가져오기
            graph_files = [f for f in os.listdir(self.br_name) if f.endswith('.png')]
            
            if graph_files:
                # 그래프를 3열로 배치하기 위한 설정
                for i in range(0, len(graph_files), 3):
                    # 첫 번째 그래프 추가
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run()
                    r.add_picture(f'{self.br_name}/{graph_files[i]}', width=Inches(4))
                    
                    # 두 번째 그래프가 있다면 추가
                    if i + 1 < len(graph_files):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run()
                        r.add_picture(f'{self.br_name}/{graph_files[i+1]}', width=Inches(4))
                    
                    # 세 번째 그래프가 있다면 추가
                    if i + 2 < len(graph_files):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run()
                        r.add_picture(f'{self.br_name}/{graph_files[i+2]}', width=Inches(4))
                    
                    # 그래프 세트 사이 간격 추가
                    doc.add_paragraph()
            else:
                doc.add_paragraph("시간이력 그래프가 없습니다.")
            
            # 문서 저장
            word_file = f'{self.br_name}/{self.br_name}_요약보고서.docx'
            doc.save(word_file)
            
            print(f'{self.br_name} 요약 보고서가 워드 문서 형식으로 생성되었습니다.')
            
        except ImportError:
            print("python-docx 라이브러리가 설치되어 있지 않습니다. 다음 명령어로 설치해주세요:")
            print("pip install python-docx")
        except Exception as e:
            print(f"워드 문서 생성 중 오류가 발생했습니다: {str(e)}")

    def calculate_weekly_reception_rate(self):
        """주별 데이터 수신율을 계산하는 함수"""
        try:
            # 주별로 그룹화하기 위해 날짜를 주 단위로 변환
            self.data_merge['Week'] = self.data_merge['계측시간'].dt.isocalendar().week
            self.data_merge['Year'] = self.data_merge['계측시간'].dt.isocalendar().year
            
            # 센서 정보 가져오기
            sensor_info = self.info[self.info['br_name'] == self.br_name]
            channel_names = sensor_info['channel_name'].tolist()
            
            # 모든 센서의 주별 수신율을 저장할 데이터프레임 생성
            weekly_rates = pd.DataFrame()
            
            # 각 센서별로 주별 수신율 계산
            for channel in channel_names:
                try:
                    avg_col = f"{channel}_AVG"
                    value_col = f"{channel}_VALUE"
                    
                    # AVG 컬럼이 없으면 VALUE 컬럼 사용
                    if avg_col not in self.data_merge.columns:
                        if value_col in self.data_merge.columns:
                            avg_col = value_col
                        else:
                            print(f"{channel} 채널의 데이터가 없습니다. 스킵합니다.")
                            continue

                    # 데이터가 숫자형인지 확인
                    if not pd.api.types.is_numeric_dtype(self.data_merge[avg_col]):
                        print(f"{channel} 채널의 데이터가 숫자형이 아닙니다. 스킵합니다.")
                        continue

                    # 주별로 그룹화하여 수신율 계산
                    weekly_data = self.data_merge.groupby(['Year', 'Week'])[avg_col].agg([
                        ('received', 'count'),  # 수신된 데이터 수
                        ('total', 'size')      # 전체 데이터 수
                    ])
                    
                    # 수신율 계산
                    weekly_data['reception_rate'] = (weekly_data['received'] / weekly_data['total'] * 100).round(1)
                    
                    # 결과를 메인 데이터프레임에 추가
                    if weekly_rates.empty:
                        weekly_rates = weekly_data[['reception_rate']].copy()
                        weekly_rates.columns = [channel]
                    else:
                        # 인덱스가 일치하는지 확인하고 병합
                        weekly_rates = weekly_rates.join(weekly_data[['reception_rate']], how='outer')
                        weekly_rates.rename(columns={'reception_rate': channel}, inplace=True)
                
                except Exception as e:
                    print(f"{self.br_name} 채널명 {channel}에서 데이터 수신율 확인 중 오류 발생: {str(e)}")
                    continue

            # 인덱스 리셋하여 Year와 Week를 컬럼으로 변환
            weekly_rates = weekly_rates.reset_index()
            
            # Year와 Week를 하나의 컬럼으로 합치기
            weekly_rates['주차'] = weekly_rates['Year'].astype(str) + '년 ' + weekly_rates['Week'].astype(str) + '주차'
            weekly_rates = weekly_rates.drop(['Year', 'Week'], axis=1)
            
            # 주차 컬럼을 첫 번째로 이동
            cols = weekly_rates.columns.tolist()
            cols = ['주차'] + [col for col in cols if col != '주차']
            weekly_rates = weekly_rates[cols]
            
            # 주별 전체 센서의 평균 수신율 계산
            weekly_avg = weekly_rates.iloc[:, 1:].mean(axis=1).round(1)
            
            # 평균 수신율을 데이터프레임에 추가
            weekly_rates['평균'] = weekly_avg
            
            # 결과를 엑셀 파일로 저장
            weekly_rates.to_excel(f'{self.br_name}/{self.br_name}_주별수신율.xlsx', index=False)
            
            weekly_rates.to_csv(f'수신율/{self.br_name}_주별수신율.csv', index=False)

            # 막대그래프 생성
            plt.figure(figsize=(15, 8))
            bars = plt.bar(weekly_rates['주차'], weekly_rates['가동율'])
            
            # 그래프 스타일 설정
            plt.ylabel('가동율 (%)', fontsize=FONT_SIZE_LABEL, fontproperties=font_prop)
            plt.ylim(0, 100)  # y축 범위 설정
            
            # x축 레이블 회전
            plt.xticks(rotation=90, ha='right', fontsize=FONT_SIZE_TICK, fontproperties=font_prop)
            
            plt.yticks(fontsize=FONT_SIZE_TICK, fontproperties=font_prop)
            # 그리드 추가
            plt.grid(True, axis='y', linestyle='--', alpha=0.7)
            
            # 각 막대 위에 값 표시
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height,
                        f'{height:.1f}%',
                        ha='center', va='bottom',
                        fontproperties=font_prop)
            
            # 그래프 외곽선 설정
            ax = plt.gca()
            for spine in ax.spines.values():
                spine.set_color('gray')
            
            # 여백 조정
            plt.tight_layout()
            
            # 그래프 저장
            plt.savefig(f'{self.br_name}/{self.br_name}_주별가동율.png', dpi=300, bbox_inches='tight')
            plt.close()
            
            print(f'{self.br_name}의 주별 데이터 수신율이 엑셀 파일로 저장되었습니다.')
            print(f'{self.br_name}의 주별 평균 수신율 그래프가 생성되었습니다.')
            
        except Exception as e:
            print(f"주별 데이터 수신율 계산 중 오류가 발생했습니다: {str(e)}")
